###################################################
#Importacion de algunas librerias como docx para
#manipular el contenido de los documentos, subprocess
#solo en linux, decimal para manejar valores de moneda
#etc.
###################################################
from docx import Document#Libraryes
import re
import subprocess
from decimal import Decimal
from functools import reduce
import tempfile

###################################################
#Datos de entrada, para propositos de funcionalidad
#en una futura version los datos vendran de una vista
#y formateados.
####################################################
"""patterns = ["<day>", "<month>", "<ye>", "<company name>", 
            "<employee name>", "<city, state>", "<state or district>"]
inputs = []"""
patterns = ["(Name)", "(CC)", "(Precio1)", "(Precio2)", "(Precio3)", "(Precio4)"]
inputs = []
temp_dir = tempfile.TemporaryDirectory(dir="temp") 

for patter in patterns:
    entry = input(f"Please type an entry for {patter}")
    inputs.append(entry)

####################################################
#Inicializacion del documento plantilla.
####################################################
doc = Document("temp/Certificado_Base.docx")

####################################################
#Ciclo donde se reemplazan algunos datos, en este caso
#nombre y cedula.
####################################################

for para in doc.paragraphs:
    for index, key in enumerate(patterns):
        pattern = re.compile(key)
        if pattern.search(para.text):
            para.text = para.text.replace(key, inputs[index])
            #print(f"Replaced {key} with {inputs[index]}!")
            
        else:
            #print("Didn't find it :/")
            pass
#print(doc.tables[0].cell(0,1).text)

####################################################
#Se reemplazan los datos de la tabla de precios.
####################################################

doc.tables[0].cell(0,1).text = doc.tables[0].cell(0,1).text.replace("(Precio1)", "${:,.2f}".format(Decimal(inputs[2])))
doc.tables[0].rows[0].cells[1].paragraphs[0].runs[0].font.bold= True
doc.tables[0].cell(1,1).text = doc.tables[0].cell(1,1).text.replace("(Precio2)", "${:,.2f}".format(Decimal(inputs[3])))
doc.tables[0].rows[1].cells[1].paragraphs[0].runs[0].font.bold= True
doc.tables[0].cell(2,1).text = doc.tables[0].cell(2,1).text.replace("(Precio3)", "${:,.2f}".format(Decimal(inputs[4])))
doc.tables[0].rows[2].cells[1].paragraphs[0].runs[0].font.bold= True
doc.tables[0].cell(3,1).text = doc.tables[0].cell(3,1).text.replace("(Precio4)", "${:,.2f}".format(Decimal(inputs[5])))
doc.tables[0].rows[3].cells[1].paragraphs[0].runs[0].font.bold= True
prices = "${:,.2f}".format(reduce(lambda price1, price2: Decimal(price1) + Decimal(price2), inputs[2:6]))
doc.tables[0].cell(4,1).text = doc.tables[0].cell(4,1).text.replace("(total)", prices)
doc.tables[0].rows[4].cells[1].paragraphs[0].runs[0].font.bold= True

####################################################
#Se guarda el documento en pdf.
####################################################
doc.save(f"{temp_dir.name}/{inputs[0]}.docx")
subprocess.call(['soffice',
                # '--headless',
                '--convert-to',
                'pdf', 
                '--outdir',
                "saves",
                f"{temp_dir.name}/{inputs[0]}.docx"])
temp_dir.cleanup();